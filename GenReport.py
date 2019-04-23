import datetime
from docx import Document
import sqlite3
import getpass
import os


class GenReport:

    @staticmethod
    def gen_report(data_base_file_name):
        # Connect To Database
        conn = sqlite3.connect(data_base_file_name)
        c = conn.cursor()
        # Extract All Database Data
        c.execute("SELECT * FROM DB")
        data = c.fetchall()
        conn.close()

        # Create Word Document
        document = Document()

        # Create Report Details Section with Date and Username Creation Information
        document.add_heading('Report Details', level=1)
        now = str(datetime.datetime.today().strftime('%Y-%m-%d'))
        document.add_paragraph('Report Created On: ' + now)
        username = getpass.getuser()
        document.add_paragraph('Report Created By: ' + username)

        # Format Network Scan Data into a table
        document.add_heading('Network Search Results', level=1)
        table = document.add_table(rows=len(data) + 1, cols=len(data[0]))
        row = table.rows[0]
        row.cells[0].bold = True
        row.cells[1].bold = True
        row.cells[0].text = 'IP Address'
        row.cells[1].text = 'Host Name'
        for i in range(1, len(data)):
            row = table.rows[i]
            row_data = data[i]
            row.cells[0].text = row_data[0]
            row.cells[1].text = row_data[1]
        table.style = "Light List Accent 1"

        # Save the document
        document.save('Network_Search_Report.docx')

        # Open Document
        os.system('start Network_Search_Report.docx')




